"""
Document text + table extraction.

Strategy per file type:
  PDF  → pdfplumber for text & lattice/ruled tables
         Tables serialised as Markdown so the LLM reads exact cell values
  DOCX → python-docx (paragraphs + native tables → Markdown)
  XLSX → openpyxl  (each sheet as Markdown table)
  TXT  → raw UTF-8

Output: list[PageContent]  — one entry per PDF page / sheet / pseudo-page
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

import pdfplumber
from docx import Document as DocxDoc
from docx.table import Table as DocxTable
import openpyxl

log = logging.getLogger(__name__)


@dataclass
class PageContent:
    page_num: int
    text: str                          # prose text
    tables: list[str] = field(default_factory=list)  # each table as Markdown string

    @property
    def full_text(self) -> str:
        """Combine prose + tables into one string for chunking."""
        parts = [self.text.strip()]
        for t in self.tables:
            parts.append(t.strip())
        return "\n\n".join(p for p in parts if p)


# ── Markdown helpers ──────────────────────────────────────────────────────────

def _rows_to_markdown(rows: list[list[str]]) -> str:
    """Convert a 2-D list of cell strings to a GFM Markdown table."""
    if not rows:
        return ""
    # Normalise row widths
    ncols = max(len(r) for r in rows)
    norm = [r + [""] * (ncols - len(r)) for r in rows]

    def _cell(v: str) -> str:
        return str(v).replace("|", "\\|").replace("\n", " ").strip()

    header = "| " + " | ".join(_cell(c) for c in norm[0]) + " |"
    sep    = "| " + " | ".join("---" for _ in norm[0]) + " |"
    body   = "\n".join(
        "| " + " | ".join(_cell(c) for c in row) + " |"
        for row in norm[1:]
    )
    return "\n".join(filter(None, [header, sep, body]))


# ── PDF extraction ────────────────────────────────────────────────────────────

def _extract_tables_from_page(page) -> list[str]:
    """
    Extract all tables on a pdfplumber page as Markdown strings.
    Uses pdfplumber's TableFinder with explicit settings for better
    detection of ruled, dashed, and colour-background tables.
    """
    tables_md = []
    try:
        finder_settings = {
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
            "snap_tolerance": 3,
            "join_tolerance": 3,
            "edge_min_length": 3,
            "min_words_vertical": 1,
            "min_words_horizontal": 1,
        }
        raw_tables = page.extract_tables(table_settings=finder_settings)
        for tbl in raw_tables:
            # Normalise: replace None cells with ""
            rows = [[str(c).strip() if c is not None else "" for c in row] for row in tbl]
            # Drop completely empty rows
            rows = [r for r in rows if any(c for c in r)]
            md = _rows_to_markdown(rows)
            if md:
                tables_md.append(md)
    except Exception as e:
        log.debug("pdfplumber table extract error: %s", e)
    return tables_md


def _text_minus_tables(page) -> str:
    """
    Return prose text with table bounding boxes blanked out
    so table content isn't double-embedded as both text and table.
    """
    try:
        found = page.find_tables()
        if not found:
            return page.extract_text() or ""
        cropped = page
        for tbl in found:
            cropped = cropped.outside_bbox(tbl.bbox)
        return cropped.extract_text() or ""
    except Exception:
        return page.extract_text() or ""


def extract_pdf(data: bytes) -> list[PageContent]:
    pages: list[PageContent] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            prose  = _text_minus_tables(page)
            tables = _extract_tables_from_page(page)
            if prose.strip() or tables:
                pages.append(PageContent(page_num=i, text=prose, tables=tables))
    return pages


# ── DOCX extraction ───────────────────────────────────────────────────────────

def _docx_table_to_markdown(tbl: DocxTable) -> str:
    rows = [[cell.text for cell in row.cells] for row in tbl.rows]
    return _rows_to_markdown(rows)


def extract_docx(data: bytes) -> list[PageContent]:
    doc = DocxDoc(io.BytesIO(data))
    paragraphs: list[str] = []
    tables: list[str] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            text = block.text_content() if hasattr(block, "text_content") else ""
            # fallback using lxml
            text = "".join(t for t in block.itertext()).strip()
            if text:
                paragraphs.append(text)
        elif tag == "tbl":
            # Find matching python-docx Table object
            from docx.oxml.ns import qn
            for tbl_obj in doc.tables:
                if tbl_obj._tbl is block:
                    tables.append(_docx_table_to_markdown(tbl_obj))
                    break

    full_text = "\n".join(paragraphs)
    # Split into pseudo-pages of ~3000 chars
    pseudo_pages: list[PageContent] = []
    step = 3000
    for i in range(max(1, (len(full_text) + step - 1) // step)):
        chunk = full_text[i * step:(i + 1) * step].strip()
        if chunk:
            pseudo_pages.append(PageContent(page_num=i + 1, text=chunk, tables=[]))

    # Attach tables to last page (they're usually appendices)
    if tables:
        if pseudo_pages:
            pseudo_pages[-1].tables = tables
        else:
            pseudo_pages.append(PageContent(page_num=1, text="", tables=tables))

    return pseudo_pages


# ── XLSX extraction ───────────────────────────────────────────────────────────

def extract_xlsx(data: bytes) -> list[PageContent]:
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    pages: list[PageContent] = []
    for i, ws in enumerate(wb.worksheets, 1):
        rows = []
        for row in ws.iter_rows(values_only=True):
            r = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in r):
                rows.append(r)
        if rows:
            md = _rows_to_markdown(rows)
            pages.append(PageContent(
                page_num=i,
                text=f"Sheet: {ws.title}",
                tables=[md] if md else [],
            ))
    return pages


# ── Plain text ────────────────────────────────────────────────────────────────

def extract_txt(data: bytes) -> list[PageContent]:
    text = data.decode("utf-8", errors="ignore")
    step = 3000
    return [
        PageContent(page_num=i + 1, text=text[i * step:(i + 1) * step])
        for i in range(max(1, (len(text) + step - 1) // step))
        if text[i * step:(i + 1) * step].strip()
    ]


# ── Dispatcher ────────────────────────────────────────────────────────────────

def extract_document(filename: str, data: bytes) -> list[PageContent]:
    fn = filename.lower()
    if fn.endswith(".pdf"):
        return extract_pdf(data)
    if fn.endswith(".docx"):
        return extract_docx(data)
    if fn.endswith((".xlsx", ".xls")):
        return extract_xlsx(data)
    return extract_txt(data)